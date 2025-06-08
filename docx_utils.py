from docx import Document
import tempfile

def save_answer_to_docx(answer: str) -> str:
    doc = Document()
    doc.add_heading("Odpowiedź Badacza", level=1)
    for line in answer.split("\n"):
        doc.add_paragraph(line)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name
